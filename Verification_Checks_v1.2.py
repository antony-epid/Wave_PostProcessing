############################################################################################################
# Verification of SUMMARY_MEANS and part_processed hourly files
# Author: CAS
# Date: 15/08/2024
# Version: 1.1 Edited to output verification checks in word document rather than in excel woork book as in version 1.0
############################################################################################################
# IMPORTING PACKAGES #
import docx
import config
import os
import pandas as pd
from docx.shared import RGBColor
import operator
from docx.enum.section import WD_ORIENTATION, WD_SECTION

# --- Creating verification log --- #
def create_verif_log(log_header):
    """
    This function creates a verification log as a docx format.
    :return: verif_log: Returning a verification log, where verification checks of the SUMMARY_MEANS and part_processed hourly files will be outputted
    """
    verif_log = docx.Document()
    verif_log.add_heading(f'{log_header} - {config.PC_DATE}')
    return verif_log

def add_header(log_header):
    """
    Adding header to verification.
    :param log_header: Header in verification log.
    :return:
    """
    verif_log.add_page_break()
    verif_log.add_heading(f'{log_header} - {config.PC_DATE}')
    save_verif_log(verif_log)

def save_verif_log(verif_log):
    """
    Function to save the verification log in the results folder.
    :param verif_log: The verification log
    :return:
    """
    verif_log_file_path = os.path.join(config.ROOT_FOLDER, config.LOG_FOLDER, f'{config.VERIF_NAME}_{config.PC_DATE}.docx')
    verif_log.save(verif_log_file_path)


# --- Adding text to log --- #
def add_text(log, text_to_log, x, y, z):
    """
    This function adds text to the verification log
    :param log: The verification log that the output gets printed to.
    :param text_to_log: The explanatory text that are being printed to the log.
    :param x, y, z: Specifies the color that the text will be printed in. Black text= (0,0,0), red text = (255,0,0), green text = (0,155,0)
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_to_log}")
    run.bold = True
    run.font.color.rgb = RGBColor(x, y, z)

def add_text_no_error(log, text_no_error):
    """
    This function adds text to the verification log if no errors and prints it in green.
    :param log: The verification log that the output gets printed to.
    :param text_no_error: The explanatory text that are being printed to the log.
    :return: None
    """
    # Adding to the verification log
    paragraph = log.add_paragraph()
    run = paragraph.add_run(f"{text_no_error}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 155, 0)


# --- Changing docx to landscape mode --- #
def landscape(log):
    """
    Changing the docx to landscape mode
    :param log: The document that is being changed.
    :return:
    """
    new_section = log.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

# --- Changing docx to portrait mode --- #
def portrait(log):
    """ Changing the docx to portrait mode
    :param log: The document that is being changed
    :return:
    """
    new_section = log.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width, new_section.page_height = new_section.page_height, new_section.page_width

# --- CHECKING IF DATASET EXISTS AND THEN READING IT IN --- #
def dataframe(file_name):
    """
    Importing dataset as dataframe and creating a flag if the dataset doesn't exist.
    :param file_name: The dataset to import.
    :return: df. The dataset as dataframe if it exists.
    :return: file_exists. Flag to indicate if the dataset exists.
    """
    dataframe_path = os.path.join(config.ROOT_FOLDER, config.RESULTS_FOLDER, config.SUMMARY_FOLDER, f'{file_name}.csv')
    if os.path.exists(dataframe_path):
        df = pd.read_csv(dataframe_path)
        file_exists = True

    if not os.path.exists(dataframe_path):
        df = None
        file_exists = False
    return df, file_exists

# --- GENERATING INCLUDE CRITERIA --- ~
def include_criteria(log, df, text_to_log, x, y, z):
    """
    This function creates an include variable, to tag if data files should be included in release based on certain criteria.
    :param log: Where the verification checks gets outputted to.
    :param df: Summary_df = The summary overall mean as a dataframe.
    :param text_to_log: The text printed to the verification log to show what is being checked.
    :param x, y, z: Specifies what color the text printed to the log should be.
    :return: df: Returning the updated dataframe.
    """
    # Generating include criteria
    df['include'] = 0
    df.loc[
        (df['Pwear'] > config.VER_PWEAR) &
        (df['Pwear'].notnull()) &
        (df['Pwear_morning'] >= config.VER_PWEAR_MORN) &
        (df['Pwear_noon'] >= config.VER_PWEAR_QUAD) &
        (df['Pwear_afternoon'] >= config.VER_PWEAR_QUAD) &
        (df['Pwear_night'] >= config.VER_PWEAR_QUAD), 'include'] = 1

    # Summarising include variable
    include_freq = df['include'].value_counts()
    include_percent = df['include'].value_counts(normalize=True) * 100
    include_cum_percent = include_percent.cumsum()

    # Adding to the verification log
    add_text(verif_log, text_to_log, x, y, z)

    # Adding table to verification log
    table = log.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Include', 'Frequency', 'Percent', 'Cumulative percent']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    # Adding the values to the table
    for idx, value in include_freq.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(value)
        row_cells[2].text = f"{include_percent[idx]}"
        row_cells[3].text = f"{include_cum_percent[idx]}"

    total_frequency = include_freq.sum()
    total_percent = include_percent.sum()

    row_cells = table.add_row().cells
    row_cells[0].text = 'Total'
    row_cells[1].text = str(total_frequency)
    row_cells[2].text = f"{total_percent}"
    row_cells[3].text = ""

    log.add_paragraph("\n")
    save_verif_log(log)

    return df

# --- Summarising startdate to check it falls inside the expected testing dates --- #
def sum_startdate(log, df, text_to_log, x, y, z):
    """
    Summarising the start time variable to be able to check if they all fall in between the expected test dates.
    :param log: Verification log where the verification checks gets outputted to.
    :param df: Summary_df = The summary overall mean as a dataframe.
    :param text_to_log: The text printed to the verification log to show what is being checked.
    :param x, y, z: Specifies the color that the text will be printed in
    :return: df: Returning the updated dataframe.
    """

    # Converting startdate to datetime format
    df['startdate'] = pd.to_datetime(df['startdate'])

    # Adding to the verification log
    add_text(log, text_to_log, x, y, z)

    # Adding table to verification log
    table = log.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Variable', 'Obs', 'Mean', 'Min', 'Max']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    # Adding values to the table
    row_cells = table.add_row().cells
    row_cells[0].text = "Start date"
    row_cells[1].text = str(df['startdate'].count())
    row_cells[2].text = str(df['startdate'].mean().strftime('%d%b%Y'))
    row_cells[3].text = str(df['startdate'].min().strftime('%d%b%Y'))
    row_cells[4].text = str(df['startdate'].max().strftime('%d%b%Y'))

    log.add_paragraph("\n")
    save_verif_log(log)

    return df

# --- Tagging duplicates --- #
def tagging_duplicates(df, variables):
    """
    This function tags duplicates in the dataframe based on specified variables
    :param df: The dataframe that the duplicates check should be performed on.
    :param variables: variables to check if there are duplicates.
    :return: df: The updated dataframe.
    """
    df['duplicates_data'] = df.duplicated(subset=variables).astype(int)
    return df

# --- Creating pwear_quad_diff variable to check that PWear is equal to the sum of Pwear of the quad times
def create_pwear_diff(df):
    """
    This function created variables in the dataframe that find the difference between pwear quadrants and wkend/wkday variables with the overall sum of these.
    :param df: The dataframe which is the imported summary means dataset.
    :return: df: The updated dataframe
    """
    df['Pwear_quad_diff'] = df['Pwear'] - (df['Pwear_morning'] + df['Pwear_noon'] + df['Pwear_afternoon'] + df['Pwear_night'])
    df['Pwear_wk_wkend_diff'] = df['Pwear'] - (df['Pwear_wkend'] + df['Pwear_wkday'])
    week_variables = ['wkend', 'wkday']

    for index, var in enumerate(week_variables):
        df[f'Pwear_{var}_quads_diff'] = df[f'Pwear_{var}'] - (df[f'Pwear_morning_{var}'] + df[f'Pwear_noon_{var}'] + df[f'Pwear_afternoon_{var}'] + df[f'Pwear_night_{var}'])
    return df

def proportion_categories(df):
    """ This function finds the differences in all variable different thresholds (so the difference between e.g., enmo_1plus and enmo_2plus and so on)
    :param df: The dataframe which is the imported summary means dataset.
    :return: The updated dataframe
    """
    if config.REMOVE_THRESHOLDS == 'No':

        for var in config.VERIFY_VARS:
            def loop_thresholds(a, b, c, df, var):
                for i in range(a, b, c):
                    l = i + c
                    df[f'{var}_prop_cat_{i}_{l}'] = df[f'{var}_{i}plus'] - df[f'{var}_{l}plus']

            loop_thresholds(0, 5, 1, summary_df, var)
            loop_thresholds(5, 150, 5, summary_df, var)
            loop_thresholds(150, 300, 10, summary_df, var)
            loop_thresholds(300, 1000, 100, summary_df, var)
            loop_thresholds(1000, 4000, 1000, summary_df, var)

            df[f'{var}_prop_cat_4000plus'] = df[f'{var}_4000plus']

            # Checking the total proportions of all categories equals the proportion of time spent above 0mg
            prop_cat_columns = [col for col in df.columns if col.startswith(f'{var}_prop_cat_')]
            df[f'{var}_total_prop'] = df[prop_cat_columns].sum(axis=1)
            df[f'{var}_total_diff'] = (df[f'{var}_total_prop'] - df[f'{var}_0plus']).abs()
    return df


# --- Printing verification checks to log --- #
def verif_checks(comparison_operator, variable, cut_off, df, log, text_to_log, column_number, list_of_headers, text_no_error):
    """
    This function compare the value of a defined variable to a cut_off value. If the variable is above/below a certain cut-off, a warning message will be printed to the data QC log and data for this file will be printed to the log.
    :param comparison_operator: Defining if the variable we are checking should be <, > or not equal to a cut-off value.
    :param variable: The variable from the dataframe that we are checking.
    :param cut_off: The cut-off that we are checking if a variable is <, > or not equal to.
    :param df: Dataframe with variables that are being checked.
    :param log: Vericiation log where the check are being outputted to.
    :param text_to_log: The text that we want to print to the verification log if some data is not how we expect (as of above conditions)
    :param column_number: The number of columns that we want printing to the table in the verification log.
    :param list_of_headers: Listing the variables that we want printed in the verification log in case the conditions above are met.
    :param x, y, z: Specifying the color that the text in the verification log is being printed in.
    :param text_no_error: The text that we want to print to the
    :return: None
    """
    # Defining mapping of string operators to their python function to call them in the function
    operators = {
        "<": operator.lt,
        ">": operator.gt,
        "!=": operator.ne,
        ">=": operator.ge,
        "<=": operator.le
    }

    # Creating variables as a list that can store multiple variables
    if isinstance(variable, str):
        variable = [variable]

    # Checking if the length of the list 'variables' is 1 or 2 and then creating the condition that needs checking based on this:
    if len(variable) == 1:
        # Counting
        condition = operators[comparison_operator](df[variable[0]], cut_off)
    elif len(variable) == 2:
        condition1 = operators[comparison_operator](df[variable[0]], cut_off)
        condition2 = df[variable[1]].notna()
        condition = condition1 & condition2
    else:
        raise ValueError("This function only support 1 or 2 values")

    count = df[condition].shape[0]
    if count > 0:

        # Adding to the verif_log
        add_text(log, text_to_log, 255, 0, 0)

        # Adding table to verif log
        table = log.add_table(rows=1, cols=column_number)
        table.style = 'Table Grid'

        # Defining the header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(list_of_headers):
            run = hdr_cells[i].paragraphs[0].add_run(header)
            run.bold = True

        # Adding the data to the table:
        filtered_df = summary_df[condition]
        for _, row in filtered_df.iterrows():
            row_cells = table.add_row().cells
            for i, header in enumerate(list_of_headers):
                row_cells[i].text = str(row[header])

        log.add_paragraph("\n")
        save_verif_log(log)

    else:
        # Adding to the qc_log
        add_text_no_error(log, text_no_error)
        log.add_paragraph("\n")
        save_verif_log(log)


# --- LOOKING AT OUTLIERS IN DATA --- #
def outliers(df, log, list_variables, extra_variable, sort, text_to_log):
    """
    This function prints out sorted variables to the verification log to check for outliers.
    :param df: Dataframe with variables that are being checked.
    :param log: Vericiation log where the check are being outputted to.
    :param list_variables: list of variables that are being printed to the log to check for outliers.
    :param extra_variable: If REMOVE_THRESHOLD is NO this extra variable will be added to the list of variables that are being printed to the verification log.
    :param sort: Before printing the data to the log, the data will be sorted by this variable so it is easier to see the lowest and highest values.
    :param text_to_log: The text that we want to print to the verification log
    :return: None
    """
    # Adding variable to list of variables to print out
    if config.REMOVE_THRESHOLDS == 'No' and extra_variable is not None:
        list_variables.append(extra_variable)

    # Adding text and table to log
    add_text(log, text_to_log, 0,0, 0)
    table = log.add_table(rows=1, cols=len(list_variables))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, var in enumerate(list_variables):
        run = hdr_cells[i].paragraphs[0].add_run(var)
        run.bold = True

    # Sorting dataframe by variable
    df = df.sort_values(by=sort)

    # Adding data to table
    for index, row_data in df.iterrows():
        if row_data[sort] != -1 and not pd.isna(row_data[sort]):
            row_cells = table.add_row().cells
            for i, var in enumerate(list_variables):
                row_cells[i].text = str(row_data[var])

    log.add_paragraph("\n")
    save_verif_log(log)


# Add summary statistic headers to log:
def sum_stat_header(log):
    """
    Creating a table and printing headers to the table
    :param log: The verification log where the table is being outputted
    :return: table: Returning the table to be able to add the statistics afterwards.
    """
    table = log.add_table(rows=1, cols=9)
    table.style = 'Table Grid'

    # Defining the row headers
    headers = ['Variable', 'Count', 'Mean', 'Std', 'Min', '25%', '50%', '75%', 'Max']
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        run = hdr_cells[idx].paragraphs[0].add_run(header)
        run.bold = True

    return table


# Add summary statistics to log:
def sum_stat_to_log(table, summary_stats, variable):
    """
    Printing the summary statistics to table. The sum_stat_header needs to be run prior to running this, to create the table
    :param table: The table returned from sum_stat_header
    :param summary_stats: The summary statistics created by summarising (describe()) one or multiple variables.
    :param variable: The variable(s) that is being summarised
    :return:
    """
    row_cells = table.add_row().cells

    # Adding overall statistics
    row_cells[0].text = variable
    row_cells[1].text = str(int(summary_stats['count']))
    row_cells[2].text = f"{summary_stats['mean']:.4f}"
    row_cells[3].text = f"{summary_stats['std']:.4f}"
    row_cells[4].text = f"{summary_stats['min']:.4f}"
    row_cells[5].text = f"{summary_stats['25%']:.4f}"
    row_cells[6].text = f"{summary_stats['50%']:.4f}"
    row_cells[7].text = f"{summary_stats['75%']:.4f}"
    row_cells[8].text = f"{summary_stats['max']:.4f}"


# --- Getting summary statistics for specified variable and printing to verification log --- #
def get_summary_stats(condition_operator, df, log, variables, text_to_log, text_no_files):
    """
    Summarising variables and printing summary statistics to log
    :param condition_operator: Condition of which the summary statistics count is being checked with
    :param df: Dataframe with the variables in
    :param log: The verification log where the statistics are being outputted
    :param variables: The variable(s) that is being summarised
    :param text_to_log: Text to print to log to describe what is being checked
    :param text_no_files: Text to print to log if no files to summarise
    :return:
    """
    if config.REMOVE_THRESHOLDS == 'No':

        operators = {
            "<": operator.lt,
            ">": operator.gt,
            "!=": operator.ne,
            ">=": operator.ge,
            "<=": operator.le
        }
        condition_func = operators.get(condition_operator)

        # Adding to the verification log
        add_text(log, text_to_log, 0, 0, 0)

        # Adding table to verification log
        table = sum_stat_header(log)

        # Flag to indicate if any files found to be included
        flag_data_include = False

        # Getting the summary statistics for each variable
        for i, variable in enumerate(variables):
            summary_stats = df[variable].describe()
            if condition_func(summary_stats['count'], 0):
                flag_data_include = True

                # Adding overall statistics
                sum_stat_to_log(table, summary_stats, variable)

        # If no files to include, print message to log
        if not flag_data_include:
            add_text(log, text_no_files, 255, 0, 0)

        log.add_paragraph("\n")
        save_verif_log(log)


# --- Comparing ENMO_N and ENMO_0_99999 --- #
def compare_enmo(df, log, create_var, var_diff, text_to_log, text_no_error):
    """
    This function is comparing ENMO_n and ENMO_0_99999
    :param df: Dataframe with variables that are being checked.
    :param log: Verficiation log where the checks are being outputted
    :param create_var: Name of variable that is created by ENMO_0_plus * 720
    :param var_diff: The difference variable
    :param text_to_log: Text to print to log to indicate what check is being done
    :param text_no_error: Text in case there are no differences
    :return:
    """
    if config.REMOVE_THRESHOLDS == 'No':
        df[create_var] = df['ENMO_0plus'] * 720

        # Generating difference between ENMO_n and ENMO_0plus_check
        df[var_diff] = df.apply(lambda x: abs(x['ENMO_n'] - x[create_var]) if x['ENMO_n'] != x[create_var] else 0, axis=1)

        # Summarising the difference
        sum_diff = df[df[var_diff] > 1][var_diff].describe()

        if sum_diff['count'] > 0:
            # Adding header to log and creating table with the summary statistics
            add_text(log, text_to_log, 0, 0, 0)
            table = sum_stat_header(log)
            sum_stat_to_log(table, sum_diff, var_diff)

        else:
            add_text_no_error(log, text_no_error)

        log.add_paragraph("\n")
        save_verif_log(log)

    return df


# --- COUNTING DIFFERENCES IN ENMO_N AND ENMO_OPLUS NOT DUE TO FIRST OR LAST TIMEPOINT IN THE FILE --- #
def diff_enmo(df, var1, var2, var3, log, text_to_log, text_no_error):
    """
    Formatting datetime variables and calculating differences between timepoint variables. Printing to log if there is any differences.
    :param df: Dataframe with the data.
    :param var1: timepoint variable --> first_file_timepoint
    :param var2: timepoint variable --> last_file_timepoint
    :param var3: timepoint variable --> DATETIME_ORIG
    :param log: Verification log.
    :param text_to_log: Text to print if any differences
    :param text_no_error: Text to print if no differences.
    :return:
    """
    # Converting timestamps to datetime format
    for var in (var1, var2, var3):
        df[var] = pd.to_datetime(df[var])

    # Generating differences between timepoint variables and DATETIME_ORIG
    for time_var in (var1, var2):
        df[f'diff_{time_var}'] = abs(df[time_var] - df[var3]) / pd.Timedelta(hours=1)

    if config.REMOVE_THRESHOLDS == 'No':
        count = df[(df['ENMO_n_0plus_diff'] > 0.01) & (df[f'diff_{var1}'] > 1) & (df[f'diff_{var2}'] > 0)].shape[0]

        if count > 0:
            add_text(log, text_to_log, 255, 0, 0)
            add_text(log, f'Count: {count}', 0, 0, 0)

        else:
            add_text_no_error(log, text_no_error)

        log.add_paragraph("\n")
        save_verif_log(log)


def sum_enmo(remove_threshold, df, log, variables, text_to_log, text_no_files):
    """
    Summarising ENMO variables and printing to log
    :param df: Dataframe with enmo variables in
    :param log: verification log
    :param variables: variable that is being summarised
    :param text_to_log: Text to explain what is being checked
    :param text_no_files: Text if there are no files to summarize
    :return:
    """

    if config.REMOVE_THRESHOLDS == remove_threshold:

        # Adding to the verification log
        add_text(log, text_to_log, 0, 0, 0)
        table = sum_stat_header(log)

        for i, variable in enumerate(variables):
            sum_stat = df[variable].describe()

            if sum_stat['count'] > 0:
                sum_stat_to_log(table, sum_stat, variable)

            else:
                add_text(log, text_no_files, 255, 0, 0)

        log.add_paragraph("\n")
        save_verif_log(log)


# --- SUMMARISING ENMO_MEAN TO CHECK FOR ANY NEGATIVE VALUES --- #
def negative_values(df, log, text_to_log, variable, list_variables, text_no_error):

    minimum = df[variable].min()
    if minimum < 0:
        add_text(log, text_to_log, 255, 0, 0)
        table = log.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # Defining the row headers
        hdr_cells = table.rows[0].cells
        for i, var in enumerate(list_variables):
            run = hdr_cells[i].paragraphs[0].add_run(var)
            run.bold = True

        # Adding data to table
        for index, row_data in df.iterrows():
            row_cells = table.add_row().cells
            for i, var in enumerate(list_variables):
                row_cells[i].text = str(row_data[var])

    # Printing text to log if "if" condition not met
    else:
        add_text_no_error(log, text_no_error)

    log.add_paragraph("\n")
    save_verif_log(log)


# --- CREATING ENMO_MEAN FLAG TO PICK OUT FILES FOR CHECKING --- #
def enmo_flag(df, log, flag, variable):
    """Counting how many files have an ENMO_mean above a specified flag
    :param df: Dataframe with the files.
    :param log: The verification log where the count it printed to.
    :param flag: The flag to count files with ENMO_mean above this
    :param variable: ENMO_mean
    :return: None
    """

    count = df[(df[variable] > flag) & (~df[variable].isna()) & (df['QC_anomalies_total'] == 0)].shape[0]
    add_text(log, text_to_log=f"Number of files with {variable} > {variable}_flag ({flag}) \n Count: {count}", x=0, y=0, z=0)
    save_verif_log(log)



if __name__ == '__main__':

    # --- SECTION 1: VERIFICATION OF OUTPUT SUMMARY OVERALL MEANS --- #
    # Creating verification log and importing summary dataframe
    verif_log = create_verif_log("VERIFICATION OF OUTPUT SUMMARY OVERALL MEANS")
    summary_df, summary_file_exists = dataframe(file_name=config.SUM_OUTPUT_FILE)

    # If dataframe exists, create include variable and print out summary of this variable as well as summarising the startdata
    if summary_file_exists:
        df = include_criteria(log=verif_log, df=summary_df, text_to_log="Summarising include variable: \n Include = 0 (include criteria are met), Include = 1 (include criteria are not met).", x=0, y=0, z=0)
        df = sum_startdate(log=verif_log, df=summary_df, text_to_log="Summarising start date: \n Check that the minimum and maximum start date falls within the expected testing dates.", x=0, y=0, z=0)
        df = tagging_duplicates(df=summary_df, variables=['device', 'generic_first_timestamp', 'generic_last_timestamp'])

        # Printing files that have not calibrated
        verif_checks(
            comparison_operator=">",
            variable="file_end_error",
            cut_off=10,
            df=summary_df,
            log=verif_log,
            text_to_log="Some files have not calibrated - check file end errors. \n Uncalibrated data will be set to missing during post-processing.",
            column_number=3,
            list_of_headers=['id', 'file_start_error', 'file_end_error'],
            text_no_error="All files have calibrated.")

        # Printing out duplicates
        verif_checks(
            comparison_operator="!=",
            variable="duplicates_data",
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are duplicates in this summary dataset. \n Add the duplicate file to the Housekeeping file to remove data from final dataset.",
            column_number=4,
            list_of_headers=['id', 'device', 'generic_first_timestamp', 'generic_last_timestamp'],
            text_no_error="There are no duplicates in this summary dataset")

        # Printing out files with negative values of ENMO_mean
        verif_checks(
            comparison_operator="<",
            variable='enmo_mean',
            cut_off=0,
            df=summary_df,
            log=verif_log,
            text_to_log="There are negative values of ENMO_mean in this summary dataset.",
            column_number=4,
            list_of_headers=['id', 'enmo_mean', 'Pwear', 'RecordLength'],
            text_no_error="There are no negative values of ENMO_mean in this summary dataset.")

        # Printing out data to look for outliers
        outliers(summary_df, verif_log, list_variables = ['id', 'enmo_mean'], extra_variable='enmo_0plus', sort='enmo_mean', text_to_log="Look through the data for potential outliers (both smallest and largest)")

        # Creating pwear difference variables and checking if the quadrants, weekend/wkday is equal to overall pwear
        df = create_pwear_diff(summary_df)
        verif_checks(
            comparison_operator=">=",
            variable=['Pwear_quad_diff', 'Pwear'],
            cut_off=0.0001,
            df=summary_df,
            log=verif_log,
            text_to_log="The sum of Pwear is not equal to the sum of Pwear for all quadrants. Look through the files below:",
            column_number=3,
            list_of_headers=['id', 'Pwear', 'Pwear_quad_diff'],
            text_no_error="Pwear = sum of Pwear for all quadrants for all IDs")
        verif_checks(
            comparison_operator=">=",
            variable=['Pwear_wk_wkend_diff', 'Pwear'],
            cut_off=0.0001,
            df=summary_df,
            log=verif_log,
            text_to_log="The sum of Pwear is not equal to the sum of Pwear for weekend and weekdays. Look through the files below:",
            column_number=3,
            list_of_headers=['id', 'Pwear', 'Pwear_wk_wkend_diff'],
            text_no_error="Pwear = sum of Pwear weekday + weekend day for all IDs")
        weekdays = ['wkday', 'wkend']
        for var in weekdays:
            verif_checks(
                comparison_operator=">=",
                variable=[f'Pwear_{var}_quads_diff', 'Pwear'],
                cut_off=0.0001,
                df=summary_df,
                log=verif_log,
                text_to_log=f"The sum of Pwear_{var} is not equal to the sum of Pwear_{var} for all quadrants. Look through the files below:",
                column_number=3,
                list_of_headers=['id', f'Pwear_{var}', f'Pwear_{var}_quads_diff'],
                text_no_error=f"Pwear_{var} = sum of Pwear_{var} for all quadrants for all IDs")

        # Creating difference variables between all thresholds and checking that the total proportion of all categories equals the proportion of time spent above 0mg
        df = proportion_categories(summary_df)
        for verif_var in config.VERIFY_VARS:
            verif_checks(
                comparison_operator=">=",
                variable=[f'{verif_var}_total_diff', 'Pwear'],
                cut_off=0.0001,
                df=summary_df,
                log=verif_log,
                text_to_log=f"List of sum of {verif_var} proportions does not equal the total proportion. Look through the files below:",
                column_number=3,
                list_of_headers=['id', f'{verif_var}_0plus', f'{verif_var}_total_prop'],
                text_no_error=f"List of sum of {verif_var} proportions is equal to the total proportion for all files.")

        # Getting summary statistics for ENMO variables (overall and only on data that meets the inclusion criteria)
        get_summary_stats(condition_operator="!=", df=summary_df, log=verif_log, variables=['enmo_0plus'], text_to_log="Overall summary statistics for enmo_0plus", text_no_files="No observations to summarize")
        get_summary_stats(condition_operator="!=", df=summary_df[summary_df['include'] == 1], log=verif_log, variables=['enmo_0plus'], text_to_log="Summary statistics for enmo_0plus where include==1", text_no_files="No observations (include == 1) to summarize")
        enmo_variables = [col for col in summary_df.columns if col.startswith('enmo_') and col.endswith('plus')]
        landscape(verif_log)
        get_summary_stats(condition_operator="!=", df=summary_df, log=verif_log, variables=enmo_variables, text_to_log="Summary statistics for enmo_*plus variables to check for any negative values", text_no_files="No observations to summarize.")
        portrait(verif_log)
        get_summary_stats(condition_operator="!=", df=summary_df[summary_df['include'] == 1], log=verif_log, variables=enmo_variables, text_to_log="Summary statistics for enmo_*plus variables where include==1, to check for any negative values ", text_no_files="No observations (include == 1) to summarize")


    # --- SECTION 2: VERIFICATION OF HOURLY FILE(S) --- #
    # Importing hourly dataframe
    hourly_df, hourly_file_exists = dataframe(file_name=config.HOUR_OUTPUT_FILE)
    add_header(log_header="VERIFICATION OF HOURLY FILES")

    # If dataframe exists, tag duplicates and print to log if there are any
    if hourly_file_exists:
        hourly_df = tagging_duplicates(
            df=hourly_df[
                (hourly_df['ENMO_mean'] != 1) &
                (hourly_df['ENMO_mean'] != 0) &
                (~hourly_df['ENMO_mean'].isnull())].copy(),
            variables=['timestamp', 'ENMO_mean', 'ENMO_sum', 'Battery_mean'])

        verif_checks(
            comparison_operator="!=",
            variable="duplicates_data",
            cut_off=0,
            df=hourly_df,
            log=verif_log,
            text_to_log="There are duplicates in this hourly dataset. \n Add the duplicate files to the Housekeeping .do file to remove data from final dataset.",
            column_number=4,
            list_of_headers=['file_id', 'timestamp', 'ENMO_mean', 'ENMO_sum', 'Battery_mean'],
            text_no_error="There are no duplicates in this hourly dataset")

        # Comparing ENMO_n and ENMO_0plus * 720. If these are not equal they will be printed to the log.
        df = compare_enmo(
            df=hourly_df,
            log=verif_log,
            create_var="ENMO_0plus_check",
            var_diff="ENMO_n_0plus_diff",
            text_to_log="Summarising the difference between ENMO_n and ENMO_0plus * 720, if these are not equal to each other",
            text_no_error="No observations to summarise: ENMO_n and ENMO_0plus * 720 are equal to each other for all observations")

        # Generating difference between the first file timepoint and DATETIME_ORIG. This can check if the ENMO_n_0plus_diff is due to the last timepoint being part way through an hour
        diff_enmo(
            df=hourly_df,
            var1='first_file_timepoint',
            var2='last_file_timepoint',
            var3='DATETIME_ORIG',
            log=verif_log,
            text_to_log="Count if diff in enmo_n & enmo_0plus NOT due to first or last timepoint in the file:",
            text_no_error="There are no observations with differences between ENMO_n and ENMO_0plus * 720")

        # Summarising ENMO_0plus if thresholds are not removes.
        sum_enmo(
            remove_threshold= "No",
            df=hourly_df,
            log=verif_log,
            variables=['ENMO_0plus'],
            text_to_log="Overall summary statistics for enmo_0plus. \n ENMO_0plus should be ~1",
            text_no_files="No observations to summarize")

        # Change verification log to landscape mode as next table is large
        landscape(verif_log)

        # Summarising all ENMO variables if thresholds are not removed
        ENMO_variables = [col for col in hourly_df.columns if col.startswith('ENMO')]
        sum_enmo(
            remove_threshold= "No",
            df=hourly_df[(df['ENMO_mean'] != -1) & (~df['ENMO_mean'].isna())],
            log=verif_log,
            variables=ENMO_variables,
            text_to_log="Summary statistics for all enmo_* variables. \n Check for any negative values",
            text_no_files="No observations to summarize")

        # Summarising ENMO_mean if thresholds are removed
        sum_enmo(
            remove_threshold= "Yes",
            df=hourly_df[(df['ENMO_mean'] != -1) & (~df['ENMO_mean'].isna())],
            log=verif_log,
            variables=['ENMO_mean'],
            text_to_log="Summary statistics for all ENMO_mean if thresholds removed. \n Check for any negative values",
            text_no_files="No observations to summarize")

        # Checking if there are any files with ENMO_mean < 0
        negative_values(df=hourly_df, log=verif_log, variable = 'ENMO_mean', text_to_log="There are negative values of ENMO_mean", list_variables=['file_id', 'ENMO_mean', 'Pwear'], text_no_error="There are no negative values of ENMO_mean")

        # Printing out data sorted by ENMO mean to look through for potential outliers
        outliers(df=hourly_df, log=verif_log, list_variables=['file_id', 'DATETIME_ORIG', 'ENMO_mean', 'ENMO_n', 'ENMO_missing', 'ENMO_sum', 'QC_anomalies_total'], extra_variable=None, sort='ENMO_mean', text_to_log="Look through the data for potential outliers (both smallest and largest)")

        # Changing verification log to portrait mode
        portrait(verif_log)

        # Creating a flag and count how many files have and ENMO_mean above the flag. The flag 600 is chosen after looking at other study data.
        enmo_flag(df=hourly_df, log=verif_log, flag=600, variable='ENMO_mean')