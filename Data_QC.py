# This file summarises the QC output results to feedback compliance information and potential device functionality issues.
# Author: CAS
# Date: 12/08/2024 (started)
# Version: 1.0 Translated from Stata code
############################################################################################################
# IMPORTING PACKAGES #
import docx
import os
import pandas as pd
import config
from docx.shared import RGBColor
import operator



# --- Creating data_qc log --- #
def create_qc_log():
    """
    This function creates a data_qc log as a docx format and saving it in the projects results folder.
    :return: qc_log
    """
    qc_log = docx.Document()
    qc_log.add_heading(f'Data QC - {config.PC_DATE}')
    qc_log.add_paragraph("\n")
    return qc_log

def save_qc_log(qc_log):
    """
    Function to save the data qc log in the results folder.
    :param qc_log:
    :return: None
    """
    qc_log_file_path = os.path.join(config.ROOT_FOLDER, config.LOG_FOLDER, f'{config.DATA_QC_LOG}.docx')
    qc_log.save(qc_log_file_path)

# --- Importing the summary file with the meta data in to check this --- #
def import_summary_dataframe():
    """
    Importing the summary means file and opening as a dataframe
    :return: summary_df: The dataframe which is the imported summary means dataset.
    """
    summary_means_file_path = os.path.join(config.ROOT_FOLDER, config.RESULTS_FOLDER, config.SUMMARY_FOLDER, f'{config.SUM_OUTPUT_FILE}.csv')
    if os.path.exists(summary_means_file_path):
        summary_df = pd.read_csv(summary_means_file_path)
    return summary_df


# --- Tagging duplicates --- #
def tagging_duplicates(summary_df):
    """
    This function tags duplicates in the dataframe based on 1. generic first timestamp and device and 2. id.
    :param summary_df: The dataframe which is the imported summary means dataset.
    :return:
    """
    summary_df['duplicates_data'] = summary_df.duplicated(subset=['generic_first_timestamp', 'device']).astype(int)
    summary_df['duplicates_id'] = summary_df.duplicated(subset=['id']).astype(int)
    return summary_df


# --- Printing file and device information to qc_lod --- #
def information_to_qc_log(qc_log, variable_to_count, text_to_log, count_mode):
    """
    This function counts the frequency of a variable in the summary dataset and prints this together with a table to the data qc log.
    :param qc_log: the qc log that it is being printed to.
    :param variable_to_count: The variable which is being counted and printed in the qc log.
    :param text_to_log: The text that will be printed above the table in the data qc log.
    :param count_mode: Defining if all should be counted or only unique values of the variable should be counted.
    :return:
    """
    # Counting frequency
    if count_mode == 'total':
        count = summary_df[variable_to_count].count()
    elif count_mode == 'unique':
        count = summary_df[variable_to_count].nunique()

    # Adding to the qc_log
    paragraph = qc_log.add_paragraph()
    run = paragraph.add_run(f"{text_to_log}: {count}")
    run.bold = True

    # Calculating frequency of the variable
    frequency = summary_df[variable_to_count].value_counts()

    # Adding table to qc log
    table = qc_log.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Defining the header row:
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Frequency'
    hdr_cells[1].text = variable_to_count

    # Adding the values to the table
    for file_value, freq in frequency.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(freq)
        row_cells[1].text = str(file_value)

    qc_log.add_paragraph("\n")
    save_qc_log(qc_log)


# --- Printing details of the start dates --- #
def checking_startdate(summary_df):
    """
    This function finds the mean, min and max of the start date and outputs it in a table in the data qc log.
    :param summary_df: The dataframe which is the imported summary means dataset.
    :return:
    """
    # Changing startdate to datetime format
    summary_df['startdate'] = pd.to_datetime(summary_df['startdate'], errors='coerce')

    # Calculating statistics
    count = summary_df['startdate'].count()
    mean = summary_df['startdate'].mean()
    min_date = summary_df['startdate'].min()
    max_date = summary_df['startdate'].max()

    # Adding to the qc_log
    paragraph = qc_log.add_paragraph()
    run = paragraph.add_run("Summary of start dates")
    run.bold = True

    # Adding table to qc log
    table = qc_log.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Printing header to the table
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Observations'
    hdr_cells[1].text = 'Mean'
    hdr_cells[2].text = 'Minimum'
    hdr_cells[3].text = 'Maximum'

    # Printing the summary details to the table
    def format_date(date):
        return date.strftime('%d%b%Y') if pd.notnull(date) else "N/A"

    row_cells = table.add_row().cells
    row_cells[0].text = str(count)
    row_cells[1].text = format_date(mean)
    row_cells[2].text = format_date(min_date)
    row_cells[3].text = format_date(max_date)

    qc_log.add_paragraph("\n")
    save_qc_log(qc_log)


# --- Printing summary statistics of all pwear variabels --- #
def pwear_statistics(qc_log, summary_df):
    """
    This function creates statistics for all Pwear variables and outputs them in the data qc log.
    :param qc_log: The log book wear the qc checks gets printed to.
    :param summary_df: The dataframe which is the imported summary means dataset.
    :return:
    """
    variables = [col for col in summary_df.columns if col.startswith('Pwear')]

    # Adding to the qc_log
    paragraph = qc_log.add_paragraph()
    run = paragraph.add_run("Compliance overview - Quadrant hours")
    run.bold = True

    # Adding table to qc log
    table = qc_log.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Defining the header row
    hdr_cells = table.rows[0].cells
    headers = ['Variable', 'Count', 'Mean', '5th percentile', '25th percentile', '50th percentile', '75th percentile', '95th percentile']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Formatting to get the statistics with only 2 decimals
    def format_statistics(stat):
        return f"{stat:.2f}" if pd.notnull(stat) else "N/A"

    # Loop through adding rows for each variable
    def add_table_row(table, data):
        row_cells = table.add_row().cells
        for i, value in enumerate(data):
            row_cells[i].text = value

    for variable in variables:
        data = summary_df[variable]
        if data.empty:
            continue

    # Adding the statistics to the table
        stats = [
            variable,
            str(data.count()),
            format_statistics(data.mean()),
            format_statistics(data.quantile(0.05)),
            format_statistics(data.quantile(0.25)),
            format_statistics(data.quantile(0.50)),
            format_statistics(data.quantile(0.75)),
            format_statistics(data.quantile(0.95))
        ]
        add_table_row(table, stats)

    qc_log.add_paragraph("\n")
    save_qc_log(qc_log)


# --- Printing qc checks to log --- #
def qc_checks(comparison_operator, variables, cut_off, text_to_log, column_number, list_of_headers):
    """
    This function compare the value of a defined variable to a cut_off value. If the variable is above/below a certain cut-off, a warning message will be printed to the data QC log and data for this file will be printed to the log.
    :param comparison_operator: Defining if the variable we are checking should be <, > or not equal to a cut-off value.
    :param variables: The variable from the dataframe that we are checking.
    :param cut_off: The cut-off that we are checking if a variable is <, > or not equal to.
    :param text_to_log: The text that we want to print to the data qc log if some data is not how we expect (as of above conditions)
    :param column_number: The number of columns that we want printing to the table in the data qc log.
    :param list_of_headers: Listing the variables that we want printed in the data qc log in case the conditions above are met.
    :return:
    """
    # Defining mapping of string operators to their python function to call them in the function
    operators = {
        "<": operator.lt,
        ">": operator.gt,
        "!=": operator.ne
    }

    # Creating variables as a list that can store multiple variables
    if isinstance(variables, str):
        variables = [variables]

    # Checking if the length of the list 'variables' is 1 or 2 and then creating the condition that needs checking based on this:
    if len(variables) == 1:
    # Counting
        condition = operators[comparison_operator](summary_df[variables[0]], cut_off)
    elif len(variables) == 2:
        condition1 = operators[comparison_operator](summary_df[variables[0]], cut_off)
        condition2 = operators[comparison_operator](summary_df[variables[1]], cut_off)
        condition = condition1 | condition2

    else:
        raise ValueError("This function only support 1 or 2 values")

    # Counting number of files that meets the condition and if this is more than 1 it enters the loop to print data to the data qc log
    count = summary_df[condition].shape[0]
    if count > 0:

        # Adding to the qc_log
        paragraph = qc_log.add_paragraph()
        run = paragraph.add_run(text_to_log)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)

        # Adding table to qc log
        table = qc_log.add_table(rows=1, cols=column_number)
        table.style = 'Table Grid'

        # Defining the header row
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(list_of_headers):
            hdr_cells[i].text = header

        # Adding the data to the table:
        filtered_df = summary_df[condition]
        for _, row in filtered_df.iterrows():
            row_cells = table.add_row().cells
            for i, header in enumerate(list_of_headers):
                row_cells[i].text = str(row[header])

        qc_log.add_paragraph("\n")
    save_qc_log(qc_log)



# --- Calling the functions --- #
if __name__ == '__main__':
    qc_log = create_qc_log()
    summary_df = import_summary_dataframe()
    tagging_duplicates(summary_df)
    information_to_qc_log(qc_log, 'id', "Number of files processed", count_mode='total')
    information_to_qc_log(qc_log, 'device', "Number of devices used", count_mode='unique')
    checking_startdate(summary_df)
    pwear_statistics(qc_log, summary_df)
    qc_checks(
        comparison_operator="<",
        variables='RecordLength',
        cut_off=config.MIN_INCLUSION_HRS,
        text_to_log=f"There are potential battery functionality issues. Some recordings are less than {config.MIN_INCLUSION_HRS} hours. It is recommended to test the following devices.",
        column_number=3,
        list_of_headers=['id', 'device', 'RecordLength'])
    qc_checks(
        comparison_operator=">",
        variables=['qc_anomaly_g', 'qc_anomaly_f'],
        cut_off=0,
        text_to_log="There are timestamp anomaly F's or G's in the files. It is recommended to remove these devices from circulation.",
        column_number=4,
        list_of_headers=['id', 'device', 'qc_anomaly_g', 'qc_anomaly_f'])
    qc_checks(
        comparison_operator="<",
        variables='qc_first_battery_pct',
        cut_off=75,
        text_to_log="Some devices were set up with low battery (>75%)",
        column_number=4,
        list_of_headers=['id', 'device', 'qc_first_battery_pct', 'qc_last_battery_pct'])
    qc_checks(
        comparison_operator="<",
        variables='qc_last_battery_pct',
        cut_off=10,
        text_to_log="Some devices had a low battery percentage at the end of recording.",
        column_number=4,
        list_of_headers=['id', 'device', 'qc_first_battery_pct', 'qc_last_battery_pct'])
    qc_checks(
        comparison_operator="!=",
        variables='frequency',
        cut_off=config.PROTOCOL_FREQUENCY,
        text_to_log=f"There are files not initialised at the protocol frequency of {config.PROTOCOL_FREQUENCY} Hz",
        column_number=3,
        list_of_headers=['id', 'device', 'frequency']
    )
    qc_checks(
        comparison_operator="!=",
        variables='duplicates_data',
        cut_off=0,
        text_to_log="There are files containing duplicate data.",
        column_number=3,
        list_of_headers=['id', 'device', 'generic_first_timestamp']
    )
    qc_checks(
        comparison_operator="!=",
        variables='duplicates_id',
        cut_off=0,
        text_to_log="There are files with duplicate IDs.",
        column_number=3,
        list_of_headers=['id', 'device', 'generic_first_timestamp']
    )
    qc_checks(
        comparison_operator=">",
        variables='file_end_error',
        cut_off=10,
        text_to_log="There are files that have not calibrated. Check the results are set to missing and feedback is not generated.",
        column_number=5,
        list_of_headers=['id', 'device', 'generic_first_timestamp', 'Pwear', 'file_end_error']

    )







